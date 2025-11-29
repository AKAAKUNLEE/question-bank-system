# -*- coding: utf-8 -*-
"""
路由模块
"""
import os
import sqlite3
import json
import random
from flask import Blueprint, render_template, redirect, url_for, request, flash, send_file, jsonify, session
from flask_login import login_required, current_user
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.auth import get_db, query_db, execute_db, admin_required
from app.utils import allowed_file, extract_questions_from_file

# 创建蓝图
main_bp = Blueprint('main', __name__)
library_bp = Blueprint('library', __name__, url_prefix='/libraries')
question_bp = Blueprint('question', __name__, url_prefix='/questions')
paper_bp = Blueprint('paper', __name__, url_prefix='/papers')


@main_bp.route('/')
def index():
    """首页路由"""
    return render_template('index.html')


@main_bp.route('/upload_template', methods=['GET'])
def upload_template():
    """下载上传模板"""
    # 创建模板文件
    doc = Document()
    doc.add_heading('题目上传模板', 0)
    doc.add_paragraph('请按照以下格式编写题目：')
    
    doc.add_heading('单选题示例：', level=1)
    doc.add_paragraph('1. 以下哪个是Python的关键字？')
    doc.add_paragraph('A. print')
    doc.add_paragraph('B. def')
    doc.add_paragraph('C. function')
    doc.add_paragraph('D. var')
    doc.add_paragraph('答案：B')
    
    doc.add_heading('多选题示例：', level=1)
    doc.add_paragraph('2. 以下哪些是Python的内置数据类型？')
    doc.add_paragraph('A. list')
    doc.add_paragraph('B. dict')
    doc.add_paragraph('C. array')
    doc.add_paragraph('D. tuple')
    doc.add_paragraph('答案：ABD')
    
    doc.add_heading('判断题示例：', level=1)
    doc.add_paragraph('3. Python是一种编译型语言。')
    doc.add_paragraph('答案：错误')
    
    doc.add_heading('简答题示例：', level=1)
    doc.add_paragraph('4. 请简述Python的特点。')
    doc.add_paragraph('答案：Python是一种解释型、面向对象、动态数据类型的高级程序设计语言。')
    
    # 保存文件到临时位置
    temp_file = os.path.join('uploads', 'template.docx')
    os.makedirs(os.path.dirname(temp_file), exist_ok=True)
    doc.save(temp_file)
    
    return send_file(temp_file, as_attachment=True, download_name='题目上传模板.docx')


@main_bp.route('/statistics')
@login_required
def statistics():
    """统计分析路由"""
    # 获取题库统计
    libraries = query_db('SELECT * FROM libraries')
    total_libraries = len(libraries)
    
    # 获取题目统计
    questions = query_db('SELECT * FROM questions')
    total_questions = len(questions)
    
    # 按题型统计
    question_types = {
        'single_choice': 0,
        'multiple_choice': 0,
        'true_false': 0,
        'short_answer': 0,
        'essay': 0
    }
    
    for q in questions:
        if q['question_type'] in question_types:
            question_types[q['question_type']] += 1
    
    # 按难度统计
    question_difficulties = {
        'easy': 0,
        'medium': 0,
        'hard': 0
    }
    
    for q in questions:
        if q['difficulty'] in question_difficulties:
            question_difficulties[q['difficulty']] += 1
    
    # 获取试卷统计
    papers = query_db('SELECT * FROM papers')
    total_papers = len(papers)
    
    return render_template('statistics.html',
                         total_libraries=total_libraries,
                         total_questions=total_questions,
                         question_types=question_types,
                         question_difficulties=question_difficulties,
                         total_papers=total_papers)


@library_bp.route('/')
@login_required
def libraries():
    """题库列表路由"""
    libraries = query_db('SELECT * FROM libraries ORDER BY created_at DESC')
    return render_template('libraries.html', libraries=libraries)


@library_bp.route('/create', methods=['POST'])
@login_required
def create_library():
    """创建题库路由"""
    name = request.form.get('name')
    description = request.form.get('description')
    
    if not name:
        flash('题库名称不能为空！', 'danger')
        return redirect(url_for('library.libraries'))
    
    # 检查题库是否已存在
    existing = query_db('SELECT * FROM libraries WHERE name = ?', [name], one=True)
    if existing:
        flash('题库名称已存在！', 'danger')
        return redirect(url_for('library.libraries'))
    
    # 创建题库
    execute_db(
        'INSERT INTO libraries (name, description, user_id) VALUES (?, ?, ?)',
        [name, description, current_user.id]
    )
    
    flash('题库创建成功！', 'success')
    return redirect(url_for('library.libraries'))


@library_bp.route('/<int:library_id>/delete')
@login_required
def delete_library(library_id):
    """删除题库路由"""
    # 检查题库是否存在
    library = query_db('SELECT * FROM libraries WHERE id = ?', [library_id], one=True)
    if not library:
        flash('题库不存在！', 'danger')
        return redirect(url_for('library.libraries'))
    
    # 检查是否有题目关联
    questions = query_db('SELECT * FROM questions WHERE library_id = ?', [library_id])
    if questions:
        # 删除关联的题目
        execute_db('DELETE FROM questions WHERE library_id = ?', [library_id])
    
    # 删除题库
    execute_db('DELETE FROM libraries WHERE id = ?', [library_id])
    
    flash('题库删除成功！', 'success')
    return redirect(url_for('library.libraries'))


@question_bp.route('/<int:library_id>')
@login_required
def questions(library_id):
    """题目列表路由"""
    library = query_db('SELECT * FROM libraries WHERE id = ?', [library_id], one=True)
    if not library:
        flash('题库不存在！', 'danger')
        return redirect(url_for('library.libraries'))
    
    questions = query_db('SELECT * FROM questions WHERE library_id = ? ORDER BY created_at DESC', [library_id])
    return render_template('questions.html', library=library, questions=questions)


@question_bp.route('/<int:library_id>/create', methods=['POST'])
@login_required
def create_question(library_id):
    """创建题目路由"""
    question_text = request.form.get('question_text')
    answer_text = request.form.get('answer_text')
    question_type = request.form.get('question_type')
    difficulty = request.form.get('difficulty')
    
    if not question_text or not answer_text:
        flash('题目和答案不能为空！', 'danger')
        return redirect(url_for('question.questions', library_id=library_id))
    
    # 创建题目
    execute_db(
        'INSERT INTO questions (library_id, question_text, answer_text, question_type, difficulty) VALUES (?, ?, ?, ?, ?)',
        [library_id, question_text, answer_text, question_type, difficulty]
    )
    
    flash('题目创建成功！', 'success')
    return redirect(url_for('question.questions', library_id=library_id))


@question_bp.route('/<int:question_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_question(question_id):
    """编辑题目路由"""
    question = query_db('SELECT * FROM questions WHERE id = ?', [question_id], one=True)
    if not question:
        flash('题目不存在！', 'danger')
        return redirect(url_for('library.libraries'))
    
    if request.method == 'POST':
        question_text = request.form.get('question_text')
        answer_text = request.form.get('answer_text')
        question_type = request.form.get('question_type')
        difficulty = request.form.get('difficulty')
        
        if not question_text or not answer_text:
            flash('题目和答案不能为空！', 'danger')
            return redirect(url_for('question.edit_question', question_id=question_id))
        
        # 更新题目
        execute_db(
            'UPDATE questions SET question_text = ?, answer_text = ?, question_type = ?, difficulty = ? WHERE id = ?',
            [question_text, answer_text, question_type, difficulty, question_id]
        )
        
        flash('题目更新成功！', 'success')
        return redirect(url_for('question.questions', library_id=question['library_id']))
    
    return render_template('edit_question.html', question=question)


@question_bp.route('/<int:question_id>/delete')
@login_required
def delete_question(question_id):
    """删除题目路由"""
    question = query_db('SELECT * FROM questions WHERE id = ?', [question_id], one=True)
    if not question:
        flash('题目不存在！', 'danger')
        return redirect(url_for('library.libraries'))
    
    execute_db('DELETE FROM questions WHERE id = ?', [question_id])
    flash('题目删除成功！', 'success')
    return redirect(url_for('question.questions', library_id=question['library_id']))


@question_bp.route('/<int:library_id>/upload', methods=['GET', 'POST'])
@login_required
def upload_questions(library_id):
    """上传题目路由"""
    library = query_db('SELECT * FROM libraries WHERE id = ?', [library_id], one=True)
    if not library:
        flash('题库不存在！', 'danger')
        return redirect(url_for('library.libraries'))
    
    if request.method == 'POST':
        # 检查是否有文件上传
        if 'file' not in request.files:
            flash('请选择文件！', 'danger')
            return redirect(request.url)
        
        file = request.files['file']
        
        # 检查文件是否为空
        if file.filename == '':
            flash('请选择文件！', 'danger')
            return redirect(request.url)
        
        # 检查文件类型
        if not allowed_file(file.filename):
            flash('不支持的文件类型！', 'danger')
            return redirect(request.url)
        
        # 保存文件
        upload_folder = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'uploads')
        os.makedirs(upload_folder, exist_ok=True)
        
        filename = file.filename
        file_path = os.path.join(upload_folder, filename)
        file.save(file_path)
        
        try:
            # 解析文件并添加题目
            questions = extract_questions_from_file(file_path)
            
            if not questions:
                flash('未解析到题目！', 'danger')
                return redirect(url_for('question.upload_questions', library_id=library_id))
            
            # 添加题目到数据库
            for q in questions:
                execute_db(
                    'INSERT INTO questions (library_id, question_text, answer_text, question_type, difficulty) VALUES (?, ?, ?, ?, ?)',
                    [library_id, q['question_text'], q['answer_text'], q['question_type'], q['difficulty']]
                )
            
            flash(f'成功导入 {len(questions)} 道题目！', 'success')
        except Exception as e:
            flash(f'导入失败：{str(e)}', 'danger')
        finally:
            # 删除上传的文件
            if os.path.exists(file_path):
                os.remove(file_path)
        
        return redirect(url_for('question.questions', library_id=library_id))
    
    return render_template('upload.html', library=library)


@question_bp.route('/<int:library_id>/batch_delete', methods=['POST'])
@login_required
def batch_delete_questions(library_id):
    """批量删除题目路由"""
    try:
        question_ids = request.json.get('question_ids', [])
        if not question_ids:
            return jsonify({'success': False, 'message': '请选择要删除的题目！'})
        
        # 转换为整数列表
        question_ids = [int(id) for id in question_ids]
        
        # 批量删除题目
        placeholders = ','.join(['?'] * len(question_ids))
        execute_db(f'DELETE FROM questions WHERE id IN ({placeholders}) AND library_id = ?', question_ids + [library_id])
        
        return jsonify({'success': True, 'message': f'成功删除 {len(question_ids)} 道题目！'})
    except Exception as e:
        return jsonify({'success': False, 'message': f'删除失败：{str(e)}'})


@paper_bp.route('/generate/<int:library_id>', methods=['GET', 'POST'])
@login_required
def generate_paper(library_id):
    """生成试卷路由"""
    library = query_db('SELECT * FROM libraries WHERE id = ?', [library_id], one=True)
    if not library:
        flash('题库不存在！', 'danger')
        return redirect(url_for('library.libraries'))
    
    if request.method == 'POST':
        # 获取表单数据
        paper_title = request.form.get('paper_title')
        paper_description = request.form.get('paper_description')
        
        # 获取各题型的数量和难度
        question_counts = {
            'single_choice': int(request.form.get('single_choice_count', 0)),
            'multiple_choice': int(request.form.get('multiple_choice_count', 0)),
            'true_false': int(request.form.get('true_false_count', 0)),
            'short_answer': int(request.form.get('short_answer_count', 0)),
            'essay': int(request.form.get('essay_count', 0))
        }
        
        question_difficulties = {
            'single_choice': request.form.get('single_choice_difficulty', 'all'),
            'multiple_choice': request.form.get('multiple_choice_difficulty', 'all'),
            'true_false': request.form.get('true_false_difficulty', 'all'),
            'short_answer': request.form.get('short_answer_difficulty', 'all'),
            'essay': request.form.get('essay_difficulty', 'all')
        }
        
        # 验证题目数量
        total_questions = sum(question_counts.values())
        if total_questions == 0:
            flash('请至少选择一道题目！', 'danger')
            return redirect(request.url)
        
        # 创建试卷
        execute_db(
            'INSERT INTO papers (title, description, user_id) VALUES (?, ?, ?)',
            [paper_title, paper_description, current_user.id]
        )
        
        # 获取试卷ID
        paper = query_db('SELECT * FROM papers WHERE title = ? AND user_id = ? ORDER BY created_at DESC', 
                        [paper_title, current_user.id], one=True)
        paper_id = paper['id']
        
        # 收集题目
        selected_questions = []
        question_order = 0
        
        for q_type, count in question_counts.items():
            if count == 0:
                continue
            
            # 构建查询
            query = 'SELECT * FROM questions WHERE library_id = ? AND question_type = ?'
            params = [library_id, q_type]
            
            # 添加难度条件
            if question_difficulties[q_type] != 'all':
                query += ' AND difficulty = ?'
                params.append(question_difficulties[q_type])
            
            # 执行查询
            available_questions = query_db(query, params)
            
            # 检查题目数量是否足够
            if len(available_questions) < count:
                flash(f'{q_type} 题目数量不足！', 'danger')
                return redirect(request.url)
            
            # 随机选择题目
            selected = random.sample(available_questions, count)
            
            # 添加到试卷
            for q in selected:
                question_order += 1
                selected_questions.append(q)
                
                # 保存试卷题目关联
                execute_db(
                    'INSERT INTO paper_questions (paper_id, question_id, question_order) VALUES (?, ?, ?)',
                    [paper_id, q['id'], question_order]
                )
        
        # 生成Word文档
        doc = Document()
        
        # 添加标题
        title = doc.add_heading(paper_title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加副标题
        subtitle = doc.add_paragraph(f'题库：{library["name"]}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph('')  # 空行
        
        # 添加题目
        current_type = None
        question_num = 1
        
        for q in selected_questions:
            # 如果题型变化，添加题型标题
            if q['question_type'] != current_type:
                current_type = q['question_type']
                
                # 题型中文名称
                type_names = {
                    'single_choice': '一、单选题',
                    'multiple_choice': '二、多选题',
                    'true_false': '三、判断题',
                    'short_answer': '四、简答题',
                    'essay': '五、论述题'
                }
                
                doc.add_heading(type_names[current_type], level=1)
                question_num = 1
            
            # 添加题目
            doc.add_paragraph(f'{question_num}. {q["question_text"]}')
            question_num += 1
        
        # 添加答案部分
        doc.add_page_break()
        doc.add_heading('参考答案', 0)
        
        current_type = None
        question_num = 1
        
        for q in selected_questions:
            # 如果题型变化，添加题型标题
            if q['question_type'] != current_type:
                current_type = q['question_type']
                
                # 题型中文名称
                type_names = {
                    'single_choice': '一、单选题答案',
                    'multiple_choice': '二、多选题答案',
                    'true_false': '三、判断题答案',
                    'short_answer': '四、简答题答案',
                    'essay': '五、论述题答案'
                }
                
                doc.add_heading(type_names[current_type], level=1)
                question_num = 1
            
            # 添加答案
            doc.add_paragraph(f'{question_num}. {q["answer_text"]}')
            question_num += 1
        
        # 保存文件
        upload_folder = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'uploads')
        os.makedirs(upload_folder, exist_ok=True)
        
        filename = f'试卷_{paper_title}.docx'
        file_path = os.path.join(upload_folder, filename)
        doc.save(file_path)
        
        # 返回文件下载
        return send_file(file_path, as_attachment=True, download_name=filename)
    
    return render_template('generate_paper.html', library=library)


@paper_bp.route('/')
@login_required
def papers():
    """试卷列表路由"""
    papers = query_db('SELECT * FROM papers ORDER BY created_at DESC')
    return render_template('papers.html', papers=papers)


@paper_bp.route('/<int:paper_id>')
@login_required
def paper(paper_id):
    """试卷详情路由"""
    paper = query_db('SELECT * FROM papers WHERE id = ?', [paper_id], one=True)
    if not paper:
        flash('试卷不存在！', 'danger')
        return redirect(url_for('paper.papers'))
    
    # 获取试卷题目
    paper_questions = query_db(
        'SELECT q.*, pq.question_order FROM questions q JOIN paper_questions pq ON q.id = pq.question_id WHERE pq.paper_id = ? ORDER BY pq.question_order', 
        [paper_id]
    )
    
    return render_template('view_paper.html', paper=paper, questions=paper_questions)